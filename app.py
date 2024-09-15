from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests

# Load the Llama 2 model
def load_llm(max_tokens, prompt_template):
    llm = CTransformers(
        model="llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens=max_tokens,  # Increase if content truncates
        temperature=0.7
    )
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    return llm_chain

# Fetching the image from Unsplash API
def get_src_original_url(query):
    access_key = "ya1qVTdp_TKvLQ5XeaNteHVtcFSgNZj7pHm0q7pYTfg"  # Replace with your Unsplash Access Key
    url = 'https://api.unsplash.com/search/photos'
    headers = {'Authorization': f'Client-ID {access_key}'}
    params = {'query': query, 'per_page': 1}

    response = requests.get(url, headers=headers, params=params)

    if response.status_code == 200:
        data = response.json()
        photos = data.get('results', [])
        if photos:
            return photos[0]['urls']['full']
        else:
            st.write("No photos found for the given query.")
            return None
    else:
        st.write(f"Error: {response.status_code}, {response.text}")
        return None

# Create a Word document with the generated article and image
def create_word_docx(user_input, paragraph, image_input):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    doc.add_heading('Image Input', level=1)

    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))  # Adjust the width as needed

    return doc

# Streamlit UI
st.set_page_config(layout="wide")

def main():
    st.title("Article Generator App using Llama 2")

    # Get user input for article and image
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")

    if user_input and image_input:
        col1, col2, col3 = st.columns([1, 2, 1])

        # Generate content using Llama 2
        with col1:
            st.subheader("Generated Content by Llama 2")
            st.write(f"Topic of the article is: {user_input}")
            st.write(f"Image of the article is: {image_input}")

            # Ensure prompt is specific and concise
            prompt_template = """You are an SEO and marketing expert. Write a detailed article on the topic: {user_input}. Keep it informative and under 800 words."""
            llm_call = load_llm(max_tokens=1200, prompt_template=prompt_template)  # Increased max_tokens to avoid truncation
            result = llm_call.run(user_input)  # Run the model

            if result:
                st.info("Your article has been generated successfully!")
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")

        # Fetch and display the image
        with col2:
            st.subheader("Fetched Image")
            image_url = get_src_original_url(image_input)
            if image_url:
                st.image(image_url)
            else:
                st.write("Failed to fetch image")

        # Create and download the Word document
        with col3:
            st.subheader("Final Article to Download")
            temp_image_path = "temp_image.jpg"
            if result and image_url:
                # Download the image for use in Word document
                img_data = requests.get(image_url).content
                with open(temp_image_path, 'wb') as img_file:
                    img_file.write(img_data)

                # Create Word document with article and image
                doc = create_word_docx(user_input, result, Image.open(temp_image_path))

                # Save the Word document to a BytesIO buffer
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                # Provide download link
                st.download_button(
                    label='Download Word Document',
                    data=doc_buffer,
                    file_name='generated_article.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            else:
                st.write("Complete both article generation and image fetch to enable the download.")

if __name__ == "__main__":
    main()
